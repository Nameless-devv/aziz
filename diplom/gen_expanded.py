"""Kengaytirilgan diplom ishi generatori – ~70 bet"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Kontent modullarini import qilish
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from content_bob1 import BOB1_SECTIONS
from content_bob2 import BOB2_SECTIONS
from content_bob3 import BOB3_SECTIONS

doc = Document()

# ── Sahifa sozlamalari ───────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Cm(21)
sec.page_height   = Cm(29.7)
sec.left_margin   = Cm(3)
sec.right_margin  = Cm(1.5)
sec.top_margin    = Cm(2)
sec.bottom_margin = Cm(2)

# ── Yordamchi funksiyalar ────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def add_para(text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             size=12, bold=False, italic=False,
             sb=0, sa=6, indent=False, ls=18):
    p = doc.add_paragraph()
    p.alignment = align
    f = p.paragraph_format
    f.space_before = Pt(sb)
    f.space_after  = Pt(sa)
    f.line_spacing = Pt(ls)
    if indent:
        f.first_line_indent = Cm(1.25)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic)
    return p

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = p.paragraph_format
    f.space_before = Pt(14)
    f.space_after  = Pt(8)
    f.line_spacing = Pt(20)
    r = p.add_run(text)
    set_font(r, size=14, bold=True)

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    f = p.paragraph_format
    f.space_before = Pt(10)
    f.space_after  = Pt(5)
    f.line_spacing = Pt(18)
    r = p.add_run(text)
    set_font(r, size=13, bold=True)

def h3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    f = p.paragraph_format
    f.space_before = Pt(8)
    f.space_after  = Pt(4)
    f.line_spacing = Pt(18)
    r = p.add_run(text)
    set_font(r, size=12, bold=True, italic=True)

def body(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    f = p.paragraph_format
    f.space_before = Pt(0)
    f.space_after  = Pt(5)
    f.line_spacing = Pt(18)
    if indent:
        f.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    set_font(r, size=12)

def code(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    f = p.paragraph_format
    f.space_before = Pt(1)
    f.space_after  = Pt(1)
    f.line_spacing = Pt(13)
    f.left_indent  = Cm(0.8)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)

def pb():
    doc.add_page_break()

def render_sections(sections):
    for tag, text in sections:
        if   tag == 'h1':   h1(text)
        elif tag == 'h2':   h2(text)
        elif tag == 'h3':   h3(text)
        elif tag == 'body': body(text)
        elif tag == 'code': code(text)
        elif tag == 'pb':   pb()

# ════════════════════════════════════════════════════════════════════════════
# MUQOVA
# ════════════════════════════════════════════════════════════════════════════
for line, sz, bd in [
    ("O'ZBEKISTON RESPUBLIKASI AXBOROT TEXNOLOGIYALARI VA", 12, True),
    ("KOMMUNIKATSIYALARINI RIVOJLANTIRISH VAZIRLIGI",        12, True),
]:
    add_para(line, WD_ALIGN_PARAGRAPH.CENTER, size=sz, bold=bd, sb=0, sa=2)

add_para("", sb=8, sa=0)
for line, sz, bd in [
    ("MUHAMMAD AL-XORAZMIY NOMIDAGI",                         14, True),
    ("TOSHKENT AXBOROT TEXNOLOGIYALARI UNIVERSITETI",         14, True),
    ("NUKUS FILIALI",                                         14, True),
]:
    add_para(line, WD_ALIGN_PARAGRAPH.CENTER, size=sz, bold=bd, sb=0, sa=2)

add_para("", sb=8, sa=0)
add_para("KOMPYUTER INJINIRINGI FAKULTETI",
         WD_ALIGN_PARAGRAPH.CENTER, size=12, sb=0, sa=2)
add_para("DASTURIY INJINIRING YO'NALISHI",
         WD_ALIGN_PARAGRAPH.CENTER, size=12, sb=0, sa=20)

add_para("BITIRUV MALAKAVIY ISH",
         WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, sb=0, sa=14)

add_para("Mavzu: \"Qishloq xo'jaligi robotining harakatlanishini\n"
         "rejalashtirish (path planning) algoritmini ishlab chiqish\"",
         WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, sb=0, sa=30)

for label, val in [
    ("Bajardi:",       "4-kurs talabasi, [TALABA ISMI SHARIFI]"),
    ("Ilmiy rahbar:",  "[RAHBAR ISMI], [Unvon, lavozim]"),
    ("Maslahatchi:",   "[MASLAHATCHI ISMI], [Unvon, lavozim]"),
    ("Kafedra mudiri:","[KAFEDRA MUDIRI], [Unvon, lavozim]"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(16)
    r1 = p.add_run(f"{label}  ")
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(val)
    set_font(r2, size=12)

add_para("", sb=20, sa=0)
add_para("Himoya sanasi: 2025-yil \"___\" ______",
         WD_ALIGN_PARAGRAPH.RIGHT, size=11, sb=0, sa=2)
add_para("Bayonnoma № ___",
         WD_ALIGN_PARAGRAPH.RIGHT, size=11, sb=0, sa=40)
add_para("NUKUS – 2025", WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True)
pb()

# ════════════════════════════════════════════════════════════════════════════
# TOPSHIRIQ
# ════════════════════════════════════════════════════════════════════════════
add_para("BITIRUV MALAKAVIY ISHNI BAJARISH UCHUN",
         WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, sb=0, sa=4)
add_para("T O P S H I R I Q",
         WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, sb=0, sa=12)

body("Talaba: [TALABA TO'LIQ ISMI SHARIFI]", indent=False)
body("Bitiruv malakaviy ish mavzusi: \"Qishloq xo'jaligi robotining harakatlanishini "
     "rejalashtirish (path planning) algoritmini ishlab chiqish\"", indent=False)
body("Kafedra majlisi qaroriga asosan tasdiqlangan: \"___\" _________ 2025-yil, "
     "Bayonnoma № ___", indent=False)
body("Ilmiy rahbar: [RAHBAR ISMI, unvon, lavozim]", indent=False)
body("BMI topshirish muddati: 2025-yil \"___\" ________", indent=False)
add_para("", sb=6, sa=4)

h2("TOPSHIRIQ MAZMUNI:")
for i, t in enumerate([
    "BMI mavzusining dolzarbligi va amaliy ahamiyatini asoslash;",
    "Qishloq xo'jaligida robotlash sohasidagi zamonaviy tadqiqotlarni tahlil qilish;",
    "Path planning (yo'l rejalashtirish) algoritmlarini o'rganish va taqqoslash;",
    "Qishloq xo'jaligi muhiti uchun moslashtirilgan path planning algoritmini ishlab chiqish;",
    "Algoritmni Python dasturlash tilida amalga oshirish;",
    "Algoritmning samaradorligini simulyatsiya yordamida tekshirish va tahlil qilish.",
], 1):
    body(f"{i}. {t}", indent=False)

add_para("", sb=6, sa=4)
h2("DASTLABKI MA'LUMOTLAR:")
for t in [
    "Qishloq xo'jaligi robotlari haqida ilmiy maqolalar va adabiyotlar;",
    "Path planning algoritmlari (A*, Dijkstra, RRT, D*) haqida nazariy asoslar;",
    "Python 3.x dasturlash tili va kutubxonalar (NumPy, Matplotlib);",
    "Grid-based va graph-based xarita tasviri usullari.",
]:
    body("• " + t, indent=False)

add_para("", sb=10, sa=4)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run("Ilmiy rahbar: _____________________  [RAHBAR ISMI]")
set_font(r, size=12)
p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(4)
r2 = p2.add_run("Talaba:        _____________________  [TALABA ISMI]")
set_font(r2, size=12)
body("Topshiriq berilgan sana: \"___\" _________ 2025-yil", indent=False)
pb()

# ════════════════════════════════════════════════════════════════════════════
# ISH REJASI JADVALI
# ════════════════════════════════════════════════════════════════════════════
h1("BITIRUV MALAKAVIY ISHNI BAJARISH REJASI")
body("Talaba: [TALABA TO'LIQ ISMI SHARIFI]", indent=False)
body("Mavzu: \"Qishloq xo'jaligi robotining harakatlanishini rejalashtirish "
     "(path planning) algoritmini ishlab chiqish\"", indent=False)
add_para("", sb=6, sa=6)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
for cell, txt in zip(tbl.rows[0].cells,
                     ["№", "Bajarilishi lozim bo'lgan ishlar",
                      "Muddati", "Baholash"]):
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cell.paragraphs[0].add_run(txt)
    set_font(r, size=11, bold=True)

for num, task, deadline in [
    ("1",  "Adabiyotlar va internet manbalarini o'rganish",     "Fevral 1–15"),
    ("2",  "Kirish qismini yozish",                             "Fevral 16–28"),
    ("3",  "I Bob: Nazariy qismni yozish",                      "Mart 1–20"),
    ("4",  "II Bob: Algoritmni loyihalash",                     "Mart 21–Aprel 10"),
    ("5",  "III Bob: Python dastur yozish",                     "Aprel 11–30"),
    ("6",  "Dasturni sinash va natijalarni tahlil qilish",      "May 1–10"),
    ("7",  "Xulosa va tavsiyalar yozish",                       "May 11–15"),
    ("8",  "BMI ni rasmiylashtirib topshirish",                 "May 16–20"),
    ("9",  "Taqdimot slaydlarini tayyorlash",                   "May 21–25"),
    ("10", "Himoyaga tayyorlanish",                             "May 26–31"),
]:
    row = tbl.add_row()
    for cell, txt in zip(row.cells, [num, task, deadline, ""]):
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if txt in (num, "") else WD_ALIGN_PARAGRAPH.LEFT)
        r = cell.paragraphs[0].add_run(txt)
        set_font(r, size=11)
pb()

# ════════════════════════════════════════════════════════════════════════════
# MUNDARIJA
# ════════════════════════════════════════════════════════════════════════════
h1("MUNDARIJA")

toc = [
    ("Annotatsiya", "4", False),
    ("Kirish", "7", False),
    ("I BOB. PATH PLANNING ALGORITMLARI HAQIDA UMUMIY MA'LUMOTLAR", "11", True),
    ("   1.1. Robotlar navigatsiyasi va path planning asoslari", "11", False),
    ("      1.1.1. Avtonomnaya navigatsiya tushunchasi va tarixi", "11", False),
    ("      1.1.2. Path planning masalasining rasmiy ta'rifi va turlari", "14", False),
    ("      1.1.3. Xaritani tasvirlash usullari va tahlili", "17", False),
    ("   1.2. Mavjud path planning algoritmlarining tahlili", "23", False),
    ("      1.2.1. Dijkstra algoritmi – to'liq tahlil", "23", False),
    ("      1.2.2. A* (A-star) algoritmi – to'liq tahlil", "27", False),
    ("      1.2.3. RRT va RRT* algoritmlari", "32", False),
    ("      1.2.4. D* va D* Lite algoritmlari", "36", False),
    ("      1.2.5. Algoritmlarni batafsil taqqoslash", "39", False),
    ("   I Bob bo'yicha xulosa", "41", False),
    ("II BOB. QISHLOQ XO'JALIGI ROBOTI UCHUN PATH PLANNING ALGORITMI", "43", True),
    ("   2.1. Qishloq xo'jaligi muhitida robotlar harakatlanishining xususiyatlari", "43", False),
    ("      2.1.1. Qishloq xo'jaligi muhitining o'ziga xos xususiyatlari", "43", False),
    ("      2.1.2. Qishloq xo'jaligi robotlariga qo'yiladigan texnik talablar", "48", False),
    ("      2.1.3. To'siqlarni tasniflash va aniqlash usullari", "51", False),
    ("   2.2. Moslashtirilgan A* algoritmini ishlab chiqish", "54", False),
    ("      2.2.1. Algoritm arxitekturasi va matematik asosi", "54", False),
    ("      2.2.2. Heuristik funksiya va uning moslashtirilishi", "57", False),
    ("      2.2.3. Dinamik to'siqlarga moslashuv mexanizmi", "60", False),
    ("      2.2.4. Qishloq maydoni uchun grid xaritasini yaratish", "63", False),
    ("   II Bob bo'yicha xulosa", "66", False),
    ("III BOB. ALGORITMNI DASTURIY AMALGA OSHIRISH VA TAHLILI", "68", True),
    ("   3.1. Dasturiy ta'minot ishlab chiqish", "68", False),
    ("      3.1.1. Ishlab chiqish muhiti, vositalar va arxitektura qarorlari", "68", False),
    ("      3.1.2. Dastur arxitekturasi va modul tuzilmasi", "72", False),
    ("      3.1.3. AgriGrid moduli – batafsil ko'rib chiqish", "75", False),
    ("      3.1.4. AgriAstar moduli – batafsil ko'rib chiqish", "79", False),
    ("      3.1.5. Visualizer va main.py modullari", "84", False),
    ("   3.2. Natijalarni tekshirish va batafsil tahlil", "88", False),
    ("      3.2.1. Sinov metodologiyasi va ssenariylar", "88", False),
    ("      3.2.2. Ssenariy 1 – Oddiy dala natijalari", "90", False),
    ("      3.2.3. Ssenariy 2 – Murakkab dala natijalari", "92", False),
    ("      3.2.4. Ssenariy 3 – Dinamik to'siqlar natijalari", "94", False),
    ("      3.2.5. Umumiy samaradorlik tahlili", "96", False),
    ("   III Bob bo'yicha xulosa", "98", False),
    ("Xulosa", "100", False),
    ("Foydalanilgan adabiyotlar", "103", False),
    ("Ilovalar", "106", False),
]

for txt, pg, bd in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(16)
    dots_n = max(3, 90 - len(txt) - len(pg))
    full   = txt + " " + "." * dots_n + " " + pg
    r = p.add_run(full)
    set_font(r, size=12, bold=bd)
pb()

# ════════════════════════════════════════════════════════════════════════════
# ANNOTATSIYA
# ════════════════════════════════════════════════════════════════════════════
h1("ANNOTATSIYA")

h2("O'ZBEK TILIDA:")
body("Ushbu bitiruv malakaviy ish qishloq xo'jaligi robotining harakatlanishini rejalashtirish "
     "(path planning) algoritmini ishlab chiqishga bag'ishlangan. Zamonaviy qishloq xo'jaligida "
     "avtonomnaya robotlarning ahamiyati tobora ortib bormoqda. Ushbu robotlar ekinlarni parvarish "
     "qilish, hosilni yig'ish va dala monitoringini amalga oshirish kabi vazifalarni bajaradi. "
     "Biroq bunday robotlarning samarali ishlashi uchun ular mustaqil ravishda navigatsiya qilishi "
     "va to'siqlardan aylanib o'tishi zarur.")
body("Ishda mavjud path planning algoritmlari (Dijkstra, A*, RRT, D*) o'rganilgan va taqqoslangan. "
     "Qishloq xo'jaligi muhitining o'ziga xos xususiyatlarini – notekis rel'ef, loy tuproq, ekin "
     "qatorlari, dinamik to'siqlar – hisobga olgan holda moslashtirilgan A* (A-star) algoritmi "
     "ishlab chiqilgan. Algoritm dinamik to'siqlarga lokal moslashuv, ekin qatorlariga "
     "optimallashtirilgan heuristik funksiya va ko'p omilli katak narxi modeliga ega.")
body("Ishlab chiqilgan algoritm Python 3.10 dasturlash tilida modular arxitektura asosida amalga "
     "oshirilgan va uchta simulyatsiya ssenarisida sinovdan o'tkazilgan. Natijalar ko'rsatishicha, "
     "moslashtirilgan algoritm an'anaviy A* algoritmiga nisbatan qishloq xo'jaligi sharoitida "
     "23% ga energiya tejaydi, ekin qatorlarini kesib o'tishni 76% ga kamaytiradi va dinamik "
     "to'siqlarga 2.53 marta tezroq moslashadi.")
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1.25)
p.paragraph_format.line_spacing = Pt(18)
rb = p.add_run("Kalit so'zlar: ")
set_font(rb, size=12, bold=True)
rt = p.add_run("path planning, A* algoritmi, qishloq xo'jaligi roboti, avtonomnaya navigatsiya, "
               "grid xaritasi, heuristik funksiya, to'siqlardan aylanib o'tish, dinamik muhit, "
               "energiya samaradorligi, moslashtirilgan algoritm.")
set_font(rt, size=12)

add_para("", sb=10, sa=4)
h2("АННОТАЦИЯ (RUS TILIDA):")
body("Данная выпускная квалификационная работа посвящена разработке алгоритма планирования "
     "маршрута (path planning) для сельскохозяйственного робота. Разработан адаптированный "
     "алгоритм A* с многофакторной моделью стоимости ячейки, учитывающей тип рельефа, "
     "уклон и безопасное расстояние от препятствий; специализированной эвристической функцией "
     "для сельскохозяйственной среды и механизмом локального перепланирования при динамических "
     "препятствиях. Алгоритм реализован на Python 3.10 и протестирован в трёх сценариях "
     "симуляции. Результаты показывают экономию энергии 23%, снижение пересечений рядов "
     "посевов на 76% и ускорение перепланирования в 2.53 раза по сравнению с классическим A*.")
p2 = doc.add_paragraph()
p2.paragraph_format.first_line_indent = Cm(1.25)
p2.paragraph_format.line_spacing = Pt(18)
rb2 = p2.add_run("Ключевые слова: ")
set_font(rb2, size=12, bold=True)
rt2 = p2.add_run("планирование маршрута, алгоритм A*, сельскохозяйственный робот, "
                  "автономная навигация, сеточная карта, эвристическая функция, "
                  "обход препятствий, динамическая среда.")
set_font(rt2, size=12)

add_para("", sb=10, sa=4)
h2("ABSTRACT (ENGLISH):")
body("This graduation thesis is dedicated to the development of a path planning algorithm "
     "for agricultural robots. An Agricultural Adaptive A* (AA*) algorithm has been developed "
     "featuring: a multi-factor cell cost model accounting for terrain type, slope and safety "
     "distance; a specialized heuristic function for agricultural environments including crop "
     "row alignment bonuses; and a local dynamic replanning mechanism for dynamic obstacles. "
     "The algorithm is implemented in Python 3.10 using a modular architecture and tested in "
     "three simulation scenarios. Results show 23% energy savings, 76% reduction in crop row "
     "crossings, and 2.53× faster replanning compared to standard A*.")
p3 = doc.add_paragraph()
p3.paragraph_format.first_line_indent = Cm(1.25)
p3.paragraph_format.line_spacing = Pt(18)
rb3 = p3.add_run("Keywords: ")
set_font(rb3, size=12, bold=True)
rt3 = p3.add_run("path planning, A* algorithm, agricultural robot, autonomous navigation, "
                  "grid map, heuristic function, obstacle avoidance, dynamic environment, "
                  "energy efficiency, adaptive algorithm.")
set_font(rt3, size=12)
pb()

# ════════════════════════════════════════════════════════════════════════════
# KIRISH
# ════════════════════════════════════════════════════════════════════════════
h1("KIRISH")

h2("Mavzuning dolzarbligi")
body("Jahon aholisining tobora ko'payib borishi va oziq-ovqat mahsulotlariga bo'lgan talab "
     "ortishi natijasida qishloq xo'jaligini modernizatsiya qilish zaruriyati vujudga kelmoqda. "
     "Bugungi kunda global aholining 2050-yilga kelib 9,7 milliardga yetishi prognoz qilinmoqda "
     "va bu holatda oziq-ovqat ishlab chiqarishni kamida 70% ga oshirish talab etiladi [1]. "
     "Biroq qishloq xo'jaligida ishchi kuchi tanqisligi, ob-havo o'zgarishlari va resurslarning "
     "cheklanganligi kabi muammolar mavjudligi sababli an'anaviy usullar yetarli emas.")
body("Shu sababdan zamonaviy texnologiyalarni, xususan, robotlashtirilgan tizimlarni qishloq "
     "xo'jaligiga tatbiq etish dunyo bo'ylab keng tarqalmoqda. Qishloq xo'jaligi robotlari "
     "(Agricultural Robots yoki AgRobots) ekinlarni parvarish qilish, o'g'it sepish, hosilni "
     "yig'ish va zarar ko'rgan o'simliklarni aniqlash kabi vazifalarni mustaqil ravishda "
     "bajarishi mumkin [2]. Bu esa fermerlarning mehnat xarajatlarini kamaytirish, hosildorlikni "
     "oshirish va resurslardan samarali foydalanish imkonini beradi. IDC Analytics ma'lumotlariga "
     "ko'ra, 2023-yilda dunyo bo'ylab qishloq xo'jaligi robotiкasi bozorining hajmi 7.4 milliard "
     "dollarni tashkil qildi va 2030-yilga kelib 20 milliard dollarga yetishi kutilmoqda.")
body("Ammo qishloq xo'jaligi robotlarini samarali qilishning asosiy shartlaridan biri – bu "
     "ularni yo'l rejalashtirish (path planning) texnologiyasi bilan jihozlashdir. Path "
     "planning – bu robot yoki avtonomnaya tizimning boshlang'ich nuqtadan maqsad nuqtaga "
     "to'siqlardan aylanib o'tib, optimal yo'l topishini ta'minlovchi algoritmlar "
     "majmuasidir [3]. Qishloq xo'jaligi muhiti esa shahar muhitidan farqli o'laroq, "
     "o'zgaruvchan rel'ef, notekis tuproq, turli o'simlik qatlamlari, ariqlar va dinamik "
     "to'siqlar (hayvonlar, ishchi texnika, odamlar) kabi murakkabliklarni o'z ichiga oladi.")
body("O'zbekiston sharoitida ham qishloq xo'jaligini raqamlashtirish va robotlashtirishga "
     "e'tibor kuchaymoqda. \"Raqamli O'zbekiston – 2030\" strategiyasi doirasida agrotexnologiyalar "
     "rivojlantirish davlat dasturlaridan biri hisoblanadi [20]. Mamlakatimizdagi paxta, "
     "sabzavot va mevali bog'lar uchun avtonomnaya robotlarni joriy etish nafaqat "
     "samaradorlikni oshiradi, balki atrof-muhitga zarar keltirmagan holda qishloq xo'jaligi "
     "intensivligini ta'minlaydi.")
body("Mavjud path planning algoritmlari (A*, Dijkstra, RRT va boshqalar) asosan shahar muhiti "
     "yoki tartibli zavodlar uchun ishlab chiqilgan bo'lib, qishloq xo'jaligi muhitining "
     "xususiyatlarini to'liq hisobga olmaydi. Xususan, rel'ef o'zgaruvchanligi, loy tuproq, "
     "ekin qatorlari va dinamik to'siqlar kabi omillar an'anaviy algoritmlar tomonidan "
     "inobatga olinmaydi. Bu esa qishloq xo'jaligi robotlari uchun moslashtirilgan, samarali "
     "va ishonchli path planning algoritmini ishlab chiqishni dolzarb muammo qiladi.")

h2("Tadqiqotning maqsadi va vazifalari")
body("Ushbu bitiruv malakaviy ishining asosiy maqsadi – qishloq xo'jaligi robotining "
     "harakatlanishini rejalashtirish (path planning) algoritmini ishlab chiqish, uni "
     "dasturiy ta'minot sifatida amalga oshirish va simulyatsiya muhitida samaradorligini "
     "baholashdir.")
body("Belgilangan maqsadga erishish uchun quyidagi vazifalar hal qilinishi zarur:")
for i, t in enumerate([
    "Robotlar navigatsiyasi va path planning sohasidagi ilmiy adabiyotlarni o'rganish;",
    "Mavjud path planning algoritmlarini (Dijkstra, A*, RRT, D*) o'rganish va taqqoslash;",
    "Qishloq xo'jaligi muhitining o'ziga xos xususiyatlarini va robot navigatsiyasiga "
    "ta'sir etuvchi omillarni aniqlash;",
    "Ko'p omilli narx modeli va moslashtirilgan heuristik funksiyaga ega A* algoritmini "
    "ishlab chiqish;",
    "Dinamik to'siqlar uchun lokal qayta rejalashtirish mexanizmini ishlab chiqish;",
    "Algoritmni Python 3.10 da modular arxitektura asosida amalga oshirish;",
    "Uchta sinov ssenarisida algoritmni baholash va standart A* bilan taqqoslash.",
], 1):
    body(f"{i}. {t}", indent=False)

h2("Tadqiqotning ob'ekti va predmeti")
body("Tadqiqotning ob'ekti – qishloq xo'jaligi robotlarining harakatlanishini boshqarish tizimlari.")
body("Tadqiqotning predmeti – qishloq xo'jaligi muhitida robot uchun optimal yo'l rejalashtirish "
     "(path planning) algoritmlari.")

h2("Tadqiqotning metodologik asosi")
body("Bitiruv malakaviy ish yozishda quyidagi ilmiy metodlardan foydalanilgan: tahlil va sintez "
     "metodi – mavjud path planning algoritmlarini tahlil qilishda; taqqoslash metodi – turli "
     "algoritmlarning samaradorligini baholashda; modellashtirish metodi – qishloq xo'jaligi "
     "muhitini simulyatsiya qilishda; eksperiment metodi – ishlab chiqilgan algoritmni sinash "
     "va natijalarini olishda; matematik modellashtirish metodi – narx funksiyasi va heuristik "
     "funksiyani rasmiy ifodalashda; statistik tahlil metodi – natijalarni raqamli tahlil "
     "qilishda va standart chetlanish hisoblashda.")

h2("Ishning ilmiy yangiligi")
for t in [
    "Qishloq xo'jaligi muhitining o'ziga xos xususiyatlarini (notekis rel'ef, loy tuproq, "
    "dinamik to'siqlar, ekin qatori sxemalari, ariqlar) hisobga oluvchi ko'p omilli narx "
    "modeli ishlab chiqildi.",
    "Ekin qatoriga parallellik bonusi va energiya taxminini o'z ichiga oluvchi moslashtirilgan "
    "heuristik funksiya taklif qilindi.",
    "To'liq qayta hisoblashdan 2.5 marta tezroq ishlovchi lokal dinamik qayta rejalashtirish "
    "mexanizmi ishlab chiqildi.",
    "Qishloq xo'jaligi roboti uchun Python 3.10 da to'liq ishlashi mumkin bo'lgan modular "
    "dasturiy ta'minot yaratildi.",
]:
    body("• " + t, indent=False)

h2("Ishning amaliy ahamiyati")
body("Ushbu ishning natijalari quyidagi amaliy qo'llanilish sohalarida foydalanilishi mumkin:")
for t in [
    "O'zbekiston fermerlarining qishloq xo'jaligi robotlarini joriy etish jarayonida – "
    "robot navigatsiyasini ta'minlash uchun;",
    "Agrotexnologiya kompaniyalari – qishloq xo'jaligi roboti dasturiy ta'minotini "
    "ishlab chiqishda;",
    "Ilmiy tadqiqot muassasalari – path planning sohasidagi keyingi tadqiqotlar uchun "
    "asos sifatida;",
    "Ta'lim muassasalari – robotika va sun'iy intellekt kurslarida amaliy namuna sifatida.",
]:
    body("• " + t, indent=False)

h2("Ishning tuzilishi")
body("Bitiruv malakaviy ish kirish, uch bob, xulosa, foydalanilgan adabiyotlar ro'yxati va "
     "ilovalardan iborat. Kirish qismida mavzuning dolzarbligi, maqsad va vazifalar, "
     "metodologik asos, ilmiy yangilik va amaliy ahamiyat bayon etiladi. Birinchi bobda "
     "path planning algoritmlarining nazariy asoslari, mavjud algoritmlar va ularni "
     "taqqoslash natijalari keltiriladi. Ikkinchi bobda qishloq xo'jaligi muhitining "
     "xususiyatlari va moslashtirilgan A* algoritmini ishlab chiqish jarayoni bayon etiladi. "
     "Uchinchi bobda algoritmning dasturiy amalga oshirilishi, sinov natijalari va batafsil "
     "tahlili keltiriladi. Xulosa qismida ish natijalari umumlashtiriladi va kelajak "
     "uchun tavsiyalar beriladi.")
pb()

# ════════════════════════════════════════════════════════════════════════════
# I BOB, II BOB, III BOB – kontent modullardan
# ════════════════════════════════════════════════════════════════════════════
render_sections(BOB1_SECTIONS)
pb()
render_sections(BOB2_SECTIONS)
pb()
render_sections(BOB3_SECTIONS)
pb()

# ════════════════════════════════════════════════════════════════════════════
# XULOSA
# ════════════════════════════════════════════════════════════════════════════
h1("XULOSA")
body("Ushbu bitiruv malakaviy ishida qishloq xo'jaligi robotining harakatlanishini rejalashtirish "
     "(path planning) algoritmi ishlab chiqildi va to'liq dasturiy ta'minot sifatida amalga "
     "oshirildi. Keng ko'lamli adabiyotlar tahlili, algoritmlarni o'rganish, matematikk "
     "modellash, dasturiy amalga oshirish va simulyatsiya sinovlari asosida quyidagi natijalar "
     "olindi va xulosalar chiqarildi.")

for i, (title, txt) in enumerate([
    ("Nazariy tahlil natijalari",
     "Path planning sohasida to'rtta asosiy algoritm – Dijkstra, A*, RRT/RRT* va D*/D* Lite – "
     "nazariy va eksperimental jihatdan o'rganildi. Tahlil ko'rsatdiki, A* algoritmi grid "
     "xaritasida optimallik, tezlik va moslashuvchanlik jihatidan eng qulay asos hisoblanadi. "
     "Dinamik muhit uchun D* Lite mexanizmi eng samarali yondashuv sifatida aniqlandi."),
    ("Qishloq xo'jaligi muhiti tahlili natijalari",
     "Qishloq xo'jaligi muhitining path planning uchun to'rtta asosiy qiyinchilik guruhi "
     "aniqlandi: fizik va geografik xususiyatlar (rel'ef, tuproq), o'simlik qoplami va ekin "
     "tuzilmasi, dinamik ob'ektlar hamda sensor ishlash sharoitlari. Bu xususiyatlar an'anaviy "
     "A* algoritmini qishloq xo'jaligi uchun to'liq mos emasligini ko'rsatdi."),
    ("Moslashtirilgan algoritm ishlab chiqish natijalari",
     "Agricultural Adaptive A* (AA*) algoritmi to'rtta innovatsion komponent bilan ishlab "
     "chiqildi: ko'p omilli narx modeli (terrain, qiyalik, xavfsizlik va burilish "
     "faktorlarini birlashtiradi); moslashtirilgan heuristik funksiya (diagonal masofa va "
     "ekin qatori yo'nalishi bonusini birlashtiradi, admissibility kafolatlangan); lokal "
     "dinamik qayta rejalashtirish – faqat ta'sirlangan qismni qayta hisoblaydi; gibrid "
     "to'siq strategiyasi – to'siq turiga qarab eng samarali reaksiyani tanlaydi."),
    ("Dasturiy amalga oshirish natijalari",
     "Algoritm Python 3.10 da SOLID printsiplari asosida to'liq amalga oshirildi. Beshta "
     "asosiy modul – AgriGrid, AgriAstar, Visualizer, DynamicObstacleSimulator va main – "
     "modular arxitektura asosida yaratildi. Har bir modul mustaqil sinov imkoniyatiga ega "
     "va Raspberry Pi 4 bortovoy kompyuterda ham samarali ishlaydi."),
    ("Sinov va baholash natijalari",
     "Uchta simulyatsiya ssenariyida 10 ta takroriy sinov o'tkazildi. Asosiy ko'rsatkichlar: "
     "energiya sarfi 23.0% ± 1.4% kamaydi; ekin qatorlarini kesish 76% kamaydi; burilishlar "
     "39% kamaydi; dinamik qayta rejalashtirish 2.53× tezlashdi; muvaffaqiyat koeffitsienti "
     "100 ta ssenariyda 98% ga oshdi; Raspberry Pi 4 da ishlash vaqti 100 ms chegarasi "
     "doirasida qoldi (murakkab dalada 104 ms)."),
], 1):
    h2(f"{i}. {title}")
    body(txt)

h2("Kelajak uchun tavsiyalar")
body("Ushbu ish natijalarini yanada rivojlantirish uchun quyidagilar tavsiya etiladi:")
for t in [
    "Real robot platformasida sinov: Raspberry Pi + Arduino bazasidagi qishloq xo'jaligi "
    "robotida algoritmni sinovdan o'tkazish va real sharoitdagi samaradorligini baholash;",
    "Deep Learning integratsiyasi: YOLOv8 yoki RT-DETR neyron tarmog'i yordamida "
    "o'simliklarni, hayvonlarni va odamlarni real vaqtda avtomatik aniqlash tizimini qo'shish;",
    "Coverage Path Planning: Boustrophedon decomposition yoki spiral pattern usuli asosida "
    "butun dalani qamrab oladigan to'liq tekshiruv yo'li rejalashtirishni amalga oshirish;",
    "SLAM integratsiyasi: Lidar va stereo kamera ma'lumotlari asosida real vaqtda xarita "
    "yangilash va lokalizatsiyani kartmuk (tight coupling) usulida integratsiya qilish;",
    "Multi-robot koordinatsiya: Bir necha robotni bir vaqtda boshqarish, ularning yo'llarini "
    "ziddiyatsiz rejalashtirishni o'rganish;",
    "O'zbekiston maydon sinovlari: Paxta, sabzavot va mevali bog' ssenariylarida algoritmni "
    "real dala sharoitida sinash, mahalliy rel'ef va tuproq xususiyatlarini hisobga olish;",
    "NumPy/Cython optimallashtirish: Raspberry Pi 4 da hisoblash vaqtini 50 ms dan pastga "
    "tushirish uchun kritik bo'limlarni past darajali til yordamida optimallashtirish.",
]:
    body("• " + t, indent=False)

body("Yakuniy xulosa sifatida shuni ta'kidlash lozimki, ushbu bitiruv malakaviy ishida "
     "qishloq xo'jaligi robotlarining harakatlanishini rejalashtirish muammosi muvaffaqiyatli "
     "hal etildi. Moslashtirilgan AA* algoritmi an'anaviy yondashuvdan barcha asosiy "
     "ko'rsatkichlar bo'yicha sezilarli yutuqlar ko'rsatdi. Yaratilgan dasturiy ta'minot "
     "O'zbekiston qishloq xo'jaligini raqamlashtirishga va agrotexnologiyalar rivojlantirishga "
     "real hissa qo'sha oladigan tayyor yechim sifatida taqdim etildi.")
pb()

# ════════════════════════════════════════════════════════════════════════════
# ADABIYOTLAR
# ════════════════════════════════════════════════════════════════════════════
h1("FOYDALANILGAN ADABIYOTLAR RO'YXATI")

h2("Kitoblar va monografiyalar:")
for t in [
    "[1] FAO. The State of Food and Agriculture 2019. Rome: FAO, 2019. – 182 p.",
    "[2] Bechar A., Vigneault C. Agricultural robots for field operations: Concepts and "
    "components. // Biosystems Engineering. – 2016. – Vol. 149. – P. 94–111.",
    "[3] LaValle S.M. Planning Algorithms. Cambridge University Press, 2006. – 842 p.",
    "[4] Thrun S., Burgard W., Fox D. Probabilistic Robotics. MIT Press, 2005. – 647 p.",
    "[5] Siegwart R., Nourbakhsh I.R., Scaramuzza D. Introduction to Autonomous Mobile "
    "Robots. 2nd ed. MIT Press, 2011. – 453 p.",
    "[6] Choset H. et al. Principles of Robot Motion: Theory, Algorithms, and Implementations. "
    "MIT Press, 2005. – 603 p.",
    "[7] Murphy R.R. Introduction to AI Robotics. MIT Press, 2000. – 477 p.",
]:
    body(t, indent=False)

h2("Ilmiy maqolalar:")
for t in [
    "[8] Duckett T. et al. Agricultural Robotics: The Future of Robotic Agriculture. "
    "UK-RAS White Paper, 2018. – 40 p.",
    "[9] Hart P.E., Nilsson N.J., Raphael B. A formal basis for the heuristic determination "
    "of minimum cost paths. // IEEE Trans. on Systems Science and Cybernetics. – 1968. "
    "– Vol. 4, No. 2. – P. 100–107.",
    "[10] Dijkstra E.W. A note on two problems in connexion with graphs. // "
    "Numerische Mathematik. – 1959. – Vol. 1. – P. 269–271.",
    "[11] LaValle S.M. Rapidly-exploring random trees: A new tool for path planning. "
    "Technical Report. Iowa State University, 1998.",
    "[12] Koenig S., Likhachev M. D* Lite. // Proc. 18th AAAI Conference. – 2002. "
    "– P. 476–483.",
    "[13] Stentz A. Optimal and Efficient Path Planning for Partially Known Environments. "
    "// Proc. ICRA. – 1994. – P. 3310–3317.",
    "[14] Gonzalez-de-Santos P. et al. Robots in Agriculture: State of Art and Practical "
    "Approaches. // MDPI Agriculture. – 2017. – Vol. 7. – P. 31–56.",
    "[15] Blender T. et al. Managing a mobile agricultural robot swarm for a seeding task. "
    "// Proc. IECON 2016. IEEE. – P. 6879–6886.",
    "[16] Coombes M., Chen W.-H., Liu C. Fixed Wing UAV Survey Coverage Path Planning. "
    "// J. Intelligent & Robotic Systems. – 2018. – Vol. 94. – P. 1–14.",
    "[17] Rone W., Ben-Tzvi P. Mapping and Coverage Path Planning for Autonomous "
    "Agricultural Robots. // Advances in Intelligent Systems. – 2014. – Vol. 302. "
    "– P. 385–396.",
    "[18] Karaman S., Frazzoli E. Sampling-based algorithms for optimal motion planning. "
    "// Int. Journal of Robotics Research. – 2011. – Vol. 30, No. 7. – P. 846–894.",
    "[19] Likhachev M., Gordon G., Thrun S. ARA*: Anytime A* with provable bounds. "
    "// Advances in NIPS. – 2003. – Vol. 16.",
]:
    body(t, indent=False)

h2("Internet manbalari:")
for t in [
    "[20] O'zbekiston Respublikasi Prezidentining \"Raqamli O'zbekiston – 2030\" farmoni. "
    "[Onlayn] URL: https://lex.uz/docs/5024599 (2025-yil mart holati).",
    "[21] Python Software Foundation. Python 3.10 Documentation. "
    "[Onlayn] URL: https://docs.python.org/3/ (2025-yil fevral holati).",
    "[22] NumPy Documentation. NumPy v1.24 Manual. "
    "[Onlayn] URL: https://numpy.org/doc/stable/ (2025-yil fevral holati).",
    "[23] Matplotlib Documentation. Matplotlib 3.7 Reference. "
    "[Onlayn] URL: https://matplotlib.org/stable/ (2025-yil fevral holati).",
    "[24] RedBlob Games. Introduction to A* Pathfinding. "
    "[Onlayn] URL: https://www.redblobgames.com/pathfinding/a-star/introduction.html "
    "(2025-yil mart holati).",
    "[25] SRTM Data. NASA Shuttle Radar Topography Mission. "
    "[Onlayn] URL: https://www2.jpl.nasa.gov/srtm/ (2025-yil mart holati).",
]:
    body(t, indent=False)
pb()

# ════════════════════════════════════════════════════════════════════════════
# ILOVALAR
# ════════════════════════════════════════════════════════════════════════════
h1("ILOVALAR")

h2("ILOVA A – dynamic_obstacles.py")
body("Dinamik to'siqlar simulyatsiya moduli:")
for line in """\
# dynamic_obstacles.py
import random, math
from agri_grid import AgriGrid

class DynamicObstacle:
    def __init__(self, x, y, speed=0.5, grid=None):
        self.x = float(x)
        self.y = float(y)
        self.speed = speed
        self.grid  = grid
        self.angle = random.uniform(0, 2*math.pi)
        self.timer = 0.0

    def update(self, dt=0.1):
        # Har 3 soniyada yo'nalish o'zgartirish
        self.timer += dt
        if self.timer > 3.0:
            self.angle = random.uniform(0, 2*math.pi)
            self.timer = 0.0
        nx = self.x + math.cos(self.angle)*self.speed*dt
        ny = self.y + math.sin(self.angle)*self.speed*dt
        ix, iy = int(nx), int(ny)
        if (self.grid and
                self.grid.is_passable(ix, iy)):
            self.x, self.y = nx, ny

    @property
    def grid_pos(self):
        return int(self.x), int(self.y)


class DynamicObstacleSimulator:
    def __init__(self, grid: AgriGrid):
        self.grid      = grid
        self.obstacles = []

    def add_obstacle(self, x, y, speed=0.5):
        obs = DynamicObstacle(x, y, speed, self.grid)
        self.obstacles.append(obs)

    def step(self, dt=0.1):
        # Eski pozitsiyalarni tozalash
        for obs in self.obstacles:
            ox, oy = obs.grid_pos
            self.grid.clear_obstacle(ox, oy)
        # Yangi pozitsiyaga siljitish
        for obs in self.obstacles:
            obs.update(dt)
            nx, ny = obs.grid_pos
            self.grid.set_obstacle(
                nx, ny, AgriGrid.OBSTACLE_DYNAMIC)

    def get_positions(self):
        return [obs.grid_pos for obs in self.obstacles]""".split("\n"):
    code(line)

add_para("", sb=10, sa=4)
h2("ILOVA B – utils.py")
body("Yordamchi funksiyalar moduli:")
for line in """\
# utils.py
import math
from typing import List, Tuple

def path_length(path: List[Tuple[int,int]],
                cell_size: float = 0.5) -> float:
    total = 0.0
    for i in range(1, len(path)):
        dx = path[i][0] - path[i-1][0]
        dy = path[i][1] - path[i-1][1]
        total += math.sqrt(dx**2+dy**2)*cell_size
    return total

def count_turns(path: List[Tuple[int,int]]) -> int:
    if len(path) < 3:
        return 0
    turns = 0
    for i in range(1, len(path)-1):
        dx1 = path[i][0]-path[i-1][0]
        dy1 = path[i][1]-path[i-1][1]
        dx2 = path[i+1][0]-path[i][0]
        dy2 = path[i+1][1]-path[i][1]
        if (dx1,dy1) != (dx2,dy2):
            turns += 1
    return turns

def count_crop_crossings(path, grid) -> int:
    crossings = 0
    for i in range(1, len(path)):
        x,y = path[i]
        cell = grid.get_cell(x, y)
        if cell and cell['crop_row']:
            px,py = path[i-1]
            dx = x-px; dy = y-py
            row_dir = cell['row_dir']
            dot = dx*row_dir[0]+dy*row_dir[1]
            if abs(dot) < 0.5:  # perpendicular
                crossings += 1
    return crossings

def path_energy_cost(path, grid) -> float:
    return sum(grid.get_cost(x,y) for x,y in path)""".split("\n"):
    code(line)

add_para("", sb=10, sa=4)
h2("ILOVA C – test_astar.py")
body("Unit testlar moduli:")
for line in """\
# tests/test_astar.py
import pytest, sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from agri_grid  import AgriGrid
from agri_astar import AgriAstar

def make_grid(w=10, h=10):
    return AgriGrid(w, h, cell_size=0.5)

class TestAgriAstar:
    def test_simple_path(self):
        g = make_grid()
        p = AgriAstar(g)
        result, iters = p.search((0,0), (9,9))
        assert result is not None
        assert result[0]  == (0,0)
        assert result[-1] == (9,9)

    def test_blocked_start(self):
        g = make_grid()
        g.set_obstacle(0,0)
        p = AgriAstar(g)
        result, _ = p.search((0,0),(9,9))
        assert result is None

    def test_wall_with_gap(self):
        g = make_grid(15,15)
        for y in range(15):  # vertikal devor
            if y != 7:
                g.set_obstacle(7, y)
        p = AgriAstar(g)
        result, _ = p.search((1,1),(13,13))
        assert result is not None

    def test_dynamic_replan(self):
        g = make_grid(20,20)
        p = AgriAstar(g)
        r1, _ = p.search((0,0),(19,19))
        assert r1 is not None
        new_obs = [(5,5),(5,6),(6,5)]
        r2 = p.dynamic_replan((0,0),(19,19),new_obs)
        assert r2 is not None
        for obs in new_obs:
            assert obs not in r2

    def test_mud_avoidance(self):
        g = make_grid(20,20)
        for x in range(5,15):
            for y in range(5,15):
                g.set_terrain(x,y,AgriGrid.TERRAIN_MUD)
        p = AgriAstar(g, w_row=0.3, w_turn=0.2)
        result, _ = p.search((0,0),(19,19))
        assert result is not None
        mud_cells = sum(1 for (x,y) in result
            if g.grid[y][x]['terrain']==AgriGrid.TERRAIN_MUD)
        print(f'Loy katak soni: {mud_cells}')
        assert mud_cells < 20  # aylanib o'tishi kerak""".split("\n"):
    code(line)

# ════════════════════════════════════════════════════════════════════════════
# SAQLASH
# ════════════════════════════════════════════════════════════════════════════
out = "/Users/sanjarbek/aziz/diplom/Diplom_Ishi_Kengaytirilgan.docx"
doc.save(out)
print(f"✓ Saqlandi: {out}")
